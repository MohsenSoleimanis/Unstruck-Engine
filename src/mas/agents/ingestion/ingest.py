"""Ingestion agent — universal document parser for any file type.

Modeled after RAG-Anything's pluggable parser architecture:
  - Auto-detects file type and routes to the right parser
  - Outputs a unified content list: List[ContentItem] with standardized fields
  - Deterministic where possible (no LLM for parsing), LLM only for repair
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any

import structlog

from mas.agents.base import BaseAgent
from mas.agents.registry import registry
from mas.schemas.results import AgentResult, ResultStatus
from mas.schemas.tasks import Task

logger = structlog.get_logger()

# Parser registry — add new parsers here for new file types
_PARSERS: dict[str, Any] = {}


def register_parser(extensions: list[str]):
    """Decorator to register a parser function for file extensions."""
    def wrapper(fn):
        for ext in extensions:
            _PARSERS[ext.lower()] = fn
        return fn
    return wrapper


@register_parser([".pdf"])
async def parse_pdf(path: Path) -> list[dict[str, Any]]:
    import fitz

    doc = fitz.open(str(path))
    items: list[dict[str, Any]] = []

    for page_num in range(len(doc)):
        page = doc[page_num]

        text = page.get_text("text")
        if text.strip():
            items.append({
                "type": "text",
                "content": text,
                "page_idx": page_num + 1,
                "source": str(path),
            })

        tables = page.find_tables()
        if tables and tables.tables:
            for t_idx, table in enumerate(tables.tables):
                try:
                    items.append({
                        "type": "table",
                        "content": table.extract(),
                        "page_idx": page_num + 1,
                        "table_index": t_idx,
                        "source": str(path),
                    })
                except Exception:
                    pass

        for img_idx, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                img_bytes = pix.tobytes("png")
                items.append({
                    "type": "image",
                    "content": img_bytes,
                    "page_idx": page_num + 1,
                    "img_index": img_idx,
                    "source": str(path),
                })
            except Exception:
                items.append({
                    "type": "image",
                    "content": None,
                    "page_idx": page_num + 1,
                    "img_index": img_idx,
                    "xref": xref,
                    "source": str(path),
                })

    doc.close()
    return items


@register_parser([".docx", ".doc"])
async def parse_docx(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    items: list[dict[str, Any]] = []

    for para in doc.paragraphs:
        if para.text.strip():
            items.append({
                "type": "text",
                "content": para.text,
                "style": para.style.name,
                "source": str(path),
            })

    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        items.append({"type": "table", "content": rows, "source": str(path)})

    return items


@register_parser([".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"])
async def parse_image(path: Path) -> list[dict[str, Any]]:
    return [{
        "type": "image",
        "content": path.read_bytes(),
        "img_path": str(path),
        "source": str(path),
    }]


@register_parser([".csv"])
async def parse_csv(path: Path) -> list[dict[str, Any]]:
    import csv
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = list(reader)
    return [{"type": "table", "content": rows, "source": str(path)}]


@register_parser([".json"])
async def parse_json(path: Path) -> list[dict[str, Any]]:
    import json
    data = json.loads(path.read_text(encoding="utf-8"))
    return [{"type": "structured", "content": data, "source": str(path)}]


@register_parser([".md", ".txt", ".html", ".xml", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".log"])
async def parse_text(path: Path) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [{"type": "text", "content": text, "source": str(path)}]


@registry.register
class IngestionAgent(BaseAgent):
    """
    Universal document ingestion — parses any file into a unified content list.

    RAG-Anything pattern: pluggable parsers, unified output format, content hashing.
    Add support for new file types by registering a parser with @register_parser.
    """

    agent_type = "ingestion"
    description = "Parses any document (PDF, DOCX, images, CSV, JSON, text) into unified content items"
    version = "0.1.0"

    async def execute(self, task: Task) -> AgentResult:
        file_path = task.context.get("file_path", "")
        if not file_path:
            return self._fail(task, "No file_path in task context")

        path = Path(file_path)
        if not path.exists():
            return self._fail(task, f"File not found: {file_path}")

        ext = path.suffix.lower()
        parser = _PARSERS.get(ext)
        if not parser:
            # Fallback: try reading as text
            try:
                text = path.read_text(encoding="utf-8", errors="replace")
                items = [{"type": "text", "content": text, "source": str(path)}]
            except Exception:
                return self._fail(task, f"No parser for extension '{ext}' and file is not readable as text")
        else:
            items = await parser(path)

        content_hash = hashlib.md5(path.read_bytes()).hexdigest()

        return AgentResult(
            task_id=task.id,
            agent_id=self.agent_id,
            agent_type=self.agent_type,
            status=ResultStatus.SUCCESS,
            output={
                "file_path": str(path),
                "file_type": ext,
                "content_hash": content_hash,
                "items": self._serialize_items(items),
                "stats": {
                    "total_items": len(items),
                    "by_type": self._count_types(items),
                    "file_size_bytes": path.stat().st_size,
                },
            },
        )

    def _serialize_items(self, items: list[dict]) -> list[dict]:
        """Make items JSON-safe (convert bytes to length indicator)."""
        safe = []
        for item in items:
            copy = dict(item)
            if isinstance(copy.get("content"), bytes):
                copy["content_bytes"] = len(copy["content"])
                copy["content"] = None  # Don't serialize raw bytes
            safe.append(copy)
        return safe

    def _count_types(self, items: list[dict]) -> dict[str, int]:
        counts: dict[str, int] = {}
        for item in items:
            t = item.get("type", "unknown")
            counts[t] = counts.get(t, 0) + 1
        return counts

    def _fail(self, task: Task, msg: str) -> AgentResult:
        return AgentResult(
            task_id=task.id,
            agent_id=self.agent_id,
            agent_type=self.agent_type,
            status=ResultStatus.FAILED,
            errors=[msg],
        )
