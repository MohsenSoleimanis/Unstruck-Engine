"""Document parser agent — handles ingestion of any document type."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Any

from langchain_core.messages import HumanMessage, SystemMessage

from mas.agents.base import BaseAgent
from mas.agents.registry import registry
from mas.schemas.results import AgentResult, ResultStatus
from mas.schemas.tasks import Task


@registry.register
class ParserAgent(BaseAgent):
    """
    Parses documents into structured content.

    Inspired by:
      - RAG-Anything's pluggable parser pattern (MinerU, Docling, PaddleOCR)
      - Protocol-intelligence v30's deterministic ingestion pipeline
      - Content separation & routing pattern

    Handles: PDF, DOCX, images, markdown, HTML.
    Outputs: Unified content list with text + multimodal items.
    """

    agent_type = "parser"
    description = "Parses documents (PDF, DOCX, images) into structured content with text and multimodal items"
    version = "0.1.0"

    async def execute(self, task: Task) -> AgentResult:
        file_path = task.context.get("file_path", "")
        parse_mode = task.context.get("mode", "auto")  # auto | text_only | full

        if not file_path:
            return AgentResult(
                task_id=task.id,
                agent_id=self.agent_id,
                agent_type=self.agent_type,
                status=ResultStatus.FAILED,
                errors=["No file_path provided in task context"],
            )

        path = Path(file_path)
        if not path.exists():
            return AgentResult(
                task_id=task.id,
                agent_id=self.agent_id,
                agent_type=self.agent_type,
                status=ResultStatus.FAILED,
                errors=[f"File not found: {file_path}"],
            )

        # Determine parser based on file type
        suffix = path.suffix.lower()
        content_items = []
        text_content = ""

        if suffix == ".pdf":
            content_items, text_content = await self._parse_pdf(path)
        elif suffix in (".docx", ".doc"):
            content_items, text_content = await self._parse_docx(path)
        elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
            content_items = [{"type": "image", "path": str(path), "page_idx": 0}]
        elif suffix in (".md", ".txt", ".html"):
            text_content = path.read_text(encoding="utf-8", errors="replace")
            content_items = [{"type": "text", "content": text_content}]
        else:
            text_content = path.read_text(encoding="utf-8", errors="replace")
            content_items = [{"type": "text", "content": text_content}]

        # Separate into text vs multimodal (RAG-Anything pattern)
        text_items = [i for i in content_items if i["type"] == "text"]
        multimodal_items = [i for i in content_items if i["type"] != "text"]

        content_hash = hashlib.md5(text_content.encode()).hexdigest()[:12]

        return AgentResult(
            task_id=task.id,
            agent_id=self.agent_id,
            agent_type=self.agent_type,
            status=ResultStatus.SUCCESS,
            output={
                "file_path": str(path),
                "file_type": suffix,
                "content_hash": content_hash,
                "text_content": text_content[:10000],
                "text_items_count": len(text_items),
                "multimodal_items_count": len(multimodal_items),
                "content_items": content_items[:50],  # Cap for context size
                "total_chars": len(text_content),
            },
        )

    async def _parse_pdf(self, path: Path) -> tuple[list[dict], str]:
        """Parse PDF using PyMuPDF (deterministic, no LLM)."""
        import fitz

        doc = fitz.open(str(path))
        items = []
        full_text = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Extract text
            text = page.get_text("text")
            if text.strip():
                items.append({"type": "text", "content": text, "page_idx": page_num + 1})
                full_text.append(f"[Page {page_num + 1}]\n{text}")

            # Extract tables (basic detection via blocks)
            tables = page.find_tables()
            if tables and tables.tables:
                for table in tables.tables:
                    try:
                        table_data = table.extract()
                        items.append({
                            "type": "table",
                            "data": table_data,
                            "page_idx": page_num + 1,
                        })
                    except Exception:
                        pass

            # Extract images
            for img_idx, img in enumerate(page.get_images(full=True)):
                items.append({
                    "type": "image",
                    "xref": img[0],
                    "page_idx": page_num + 1,
                    "img_index": img_idx,
                })

        doc.close()
        return items, "\n".join(full_text)

    async def _parse_docx(self, path: Path) -> tuple[list[dict], str]:
        """Parse DOCX using python-docx."""
        from docx import Document

        doc = Document(str(path))
        items = []
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                items.append({"type": "text", "content": para.text, "style": para.style.name})
                full_text.append(para.text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            items.append({"type": "table", "data": rows})

        return items, "\n".join(full_text)
